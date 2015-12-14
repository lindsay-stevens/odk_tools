__author__ = 'Lstevens'

import os
import errno
import argparse
from collections import OrderedDict
from xlrd import open_workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

"""
Creates image files with xlsform question text.

Run on command line with "-f" parameter which is path to xlsform file.
Image settings sheet must have all parameters set. Can process multiple langs.
Places images into "out" subfolder of xlsform file directory.
Requires PIL and xlrd, written with python 2.7.6; updated to work with 
python 3.4 by changing xrange to range ("seems to work", not tested).
"""

def sheet_to_list_of_ordereddict(sheet):
    """
    Convert a xlrd Sheet to a list of OrderedDict objects.

    Dict keys are determined by the first row of values in the sheet.

    :param sheet: A xlrd Sheet object
    :return: A list of OrderedDicts
    """
    return [OrderedDict(zip(sheet._cell_values[0], row_data))
            for row_data in sheet._cell_values[1:]]

def combine_survey_choices(survey, choices):
    """
    Combine the survey sheet values with choices, if the item is a select type.

    :param survey: survey sheet values (list of OrderedDict)
    :return: survey sheet values (list of OrderedDict)
    """
    for row in survey:
        if row['type'].startswith('select'):
            list_name = row['type'].split(' ')[1]
            row['choices'] = [x for x in choices if x['list_name'] == list_name]
    return survey

def get_language_list(survey):
    """
    Get a list of unique languages specified in the survey.

    XLSForm method of defining language specific elements: 'label::en'.
    Image processing method of defining language specific elements: 'label#en'

    Assumes that single language forms still specify the language in this way.

    :param survey: survey sheet values (list of OrderedDict)
    :return: list of language name strings
    """
    headers = set()
    for header in survey._cell_values[0]:
        if '#' in header:
            headers.add(header.split('#')[1])
        if '::' in header:
            headers.add(header.split('::')[1])
    return list(headers)

def write_language_to_markdown(filepath, survey, language):
    """
    Write each xform language out as a markdown file, with choices if any

    :param survey: survey sheet values (list of OrderedDict)
    :param language: name of language to write
    :return: Nothing
    """
    file = os.path.join(os.path.dirname(filepath), '{0}.txt'.format(language))
    label_survey = 'label#{0}'.format(language)
    hint = 'hint#{0}'.format(language)
    label_choice = 'label::{0}'.format(language)
    with open(file, 'w', encoding='utf-8') as text:
        text.write('# XForm Definition')
        text.write('\n\n![](simplify_logo.png)')
        for row in survey:
            if row[label_survey] != '' and 'group' not in row['type']:
                # Header 2 if it's a group intro page with read_only text
                if 'y' in row['read_only'].lower():
                    text.write('\n\n## {0}  '.format(row[label_survey]))
                    if row[hint] != '':
                        text.write('\n_{0}_  '.format(row[hint]))
                # Bold if it is any other kind of question
                else:
                    text.write('\n\n**{0}**  '.format(row[label_survey]))
                    if row[hint] != '':
                        text.write('\n_{0}_  '.format(row[hint]))
                # If there is a choice list then write them out with checkboxes
                if row.get('choices') is not None:
                    choice_fmt = '\n- \u2610 {0}  '
                    for choice in row['choices']:
                        text.write(choice_fmt.format(choice[label_choice]))
                # If it's not a read_only, then write the type
                elif 'y' not in row['read_only'].lower():
                    text.write('\n{0}:  '.format(row['type']))


def write_language_to_docx(filepath, survey, language, version):
    """
    Write each xform language out as a markdown file, with choices if any

    :param survey: survey sheet values (list of OrderedDict)
    :param language: name of language to write
    :return: Nothing
    """
    file = os.path.join(
        os.path.dirname(filepath), '{0}_{1}.docx'.format(version, language))
    label_survey = 'label#{0}'.format(language)
    hint = 'hint#{0}'.format(language)
    label_choice = 'label::{0}'.format(language)

    doc = Document()

    # Add the logo to the top of the first page, centered
    # TODO: get the logo filename from the XLSForm instead of hardcode
    logo = os.path.join(os.path.dirname(filepath), 'simplify_logo.png')
    doc.add_picture(logo, width=Cm(8))
    img = doc.paragraphs[-1]
    img.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the document name heading
    h1 = doc.add_paragraph()
    h1.style = 'Heading 1'
    # TODO: get the form name from the XLSForm instead of hardcode
    h1.add_run('SIMPLIFY Questionnaires')
    h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set the default document styles
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(12)

    for row in survey:
        # Header 2 if it's a group intro page with read_only text
        if 'y' in row['read_only'].lower():
            # If the previous row was read_only then group hints together.
            if 'y' not in survey[survey.index(row)-1]['read_only'].lower():
                h2 = doc.add_paragraph()
                h2.style = 'Heading 2'
                h2.add_run(row[label_survey])
                h2.paragraph_format.space_before = Pt(36)
            if row[hint] != '':
                doc.add_paragraph().add_run(row[hint]).italic = True

        # Bold if it is any other kind of question
        else:
            if row['relevant'] != '':
                skip_fmt = 'Only answer if: {0}'
                skip = doc.add_paragraph()
                skip.add_run(skip_fmt.format(row['relevant'])).italic = True
                skip.paragraph_format.keep_with_next = True
                skip.paragraph_format.space_before = Pt(24)
                skip.paragraph_format.space_after = Pt(0)
            q = doc.add_paragraph()
            q_num = ''
            if row[version] != 'x':
                q_num = '{0}. '.format(int(row[version]))
            q_text = '{0}{1}'.format(q_num, row[label_survey])
            q.add_run(q_text).bold = True
            if row['relevant'] == '':
                q.paragraph_format.space_before = Pt(24)
            q.paragraph_format.keep_with_next = True

            if row[hint] != '':
                h = doc.add_paragraph()
                h.add_run(row[hint]).italic = True
                h.paragraph_format.keep_with_next = True

        # If there is a choice list then write them out with checkboxes
        if row.get('choices') is not None:
            choice_fmt = '\u2610\t{0}'
            for choice in row['choices']:
                p = doc.add_paragraph(choice_fmt.format(choice[label_choice]))
                p.paragraph_format.first_line_indent = Pt(-24)
                p.paragraph_format.left_indent = Pt(36)

                # If it's not the last choice in the list, keep it together.
                if choice['name'] != row['choices'][-1]['name']:
                    p.paragraph_format.keep_with_next = True
                    p.paragraph_format.space_after = Pt(6)

        # If it's not a read_only, then write the type
        elif 'y' not in row['read_only'].lower():
            answer_holder = 'Answer: ________________________________'
            ans = doc.add_paragraph()
            ans.add_run(answer_holder)
            ans.paragraph_format.space_before = Pt(24)

    doc.save(file)

def filter_survey_rows_for_writing(survey, languages, version):
    """
    Filter the survey item rows ready to be written out.

    For each language:
    - Only include rows which have a non-blank label
    - And which are not group begin/end rows

    :param survey: survey sheet values (list of OrderedDict)
    :param languages: list of languages in the survey
    :param version: the sub-version to filter for ('scr' or 'fu')
    :return: dictionary of filtered rows, {language: filtered survey rows list}
    """
    write_dict = dict()
    for l in languages:
        label = 'label#{0}'.format(l)
        filter_rows = [
            r for r in survey if r[label] != '' and 'group' not in r['type']]
        write_dict[l] = [r for r in filter_rows if r[version] != '']
    return write_dict

def read_xlsform(filepath):
    """
    Read the xlsform file into OrderedDicts

    :param filepath:
    :return:
    """
    workbook = open_workbook(filename=filepath)
    survey_raw = workbook.sheet_by_name('survey')
    choices_raw = workbook.sheet_by_name('choices')
    survey_dict = sheet_to_list_of_ordereddict(survey_raw)
    choices_dict = sheet_to_list_of_ordereddict(choices_raw)
    survey = combine_survey_choices(survey_dict, choices_dict)
    languages = get_language_list(survey_raw)
    filter_scr = 'paper_scr'
    filter_fu = 'paper_fu'
    survey_scr = filter_survey_rows_for_writing(survey, languages, filter_scr)
    survey_fu = filter_survey_rows_for_writing(survey, languages, filter_fu)
    surveys_to_write = {filter_scr: survey_scr, filter_fu: survey_fu}
    for v, s in surveys_to_write.items():
        for lang, rows in s.items():
            write_language_to_docx(filepath, rows, lang, v)

    #for l in languages:
    #    label = 'label#{0}'.format(l)
    #    filter_rows = [
    #        r for r in survey if r[label] != '' and 'group' not in r['type']]
    #    screening = [r for r in filter_rows if r['paper_scr'] != '']
    #    follow_up = [r for r in filter_rows if r['paper_fu'] != '']
    #    write_language_to_docx(filepath, screening, l)


def read_xlsform2(filepath):
    """
    Read form config from xls form file, write images to 'out' sub-folder.

    Requires image_settings sheet with all configs set for each language
    Looks for survey sheet columns with same language, for example:
        image_settings sheet, text_label_column = label (for value::english)
        survey sheet, label#english is used for label text (for english)
    Images are named using item name and language, like myitem_english
    """

    # open the xlsform and read the image_settings
    xls_workbook = open_workbook(filename=filepath)
    xls_image_settings = xls_workbook.sheet_by_name(sheet_name='image_settings')

    image_settings_langs = {}
    for col_index in range(1, xls_image_settings.ncols):
        col_head_value = xls_image_settings.cell_value(0, col_index)
        if not col_head_value.find('::') == -1:
            image_settings_langs[col_index] = col_head_value.split('::')[1]

    for image_settings_lang in image_settings_langs:
        language = image_settings_langs[image_settings_lang]
        image_settings = {}
        for row_index in range(1, xls_image_settings.nrows):
            image_settings[xls_image_settings.cell_value(row_index, 0)]\
                = xls_image_settings.cell_value(row_index, image_settings_lang)

        # convert the languages and type_ignore_list strings to python lists
        image_settings['type_ignore_list']\
            = image_settings['type_ignore_list'].split(',')

        # read survey sheet into a list of dicts
        xls_survey = xls_workbook.sheet_by_name(sheet_name='survey')
        xls_survey_keys = [xls_survey.cell(0, col_index).value
                           for col_index in range(xls_survey.ncols)]
        xls_survey_rows = []
        for row_index in range(1, xls_survey.nrows):
            xls_survey_row = {xls_survey_keys[col_index]:
                              xls_survey.cell(row_index, col_index).value
                              for col_index in range(xls_survey.ncols)}
            xls_survey_rows.append(xls_survey_row)

        # for each language, write images for each item
        for xls_survey_item in xls_survey_rows:
            if xls_survey_item['type'] not in \
                    image_settings['type_ignore_list']:
                image_settings_kwargs = image_settings.copy()
                image_settings_kwargs['file_path'] = os.path.dirname(filepath)
                image_settings_kwargs['file_name'] = '{0}_{1}'.format(
                    xls_survey_item[image_settings_kwargs['file_name_column']],
                    language
                )
                if len(image_settings_kwargs['nest_image_column']):
                    nest_image_path = xls_survey_item['{0}#{1}'.format(
                        image_settings_kwargs['nest_image_column'], language)]
                    if len(nest_image_path) > 0:
                        image_settings_kwargs['nest_image_path'] = \
                            nest_image_path


def create_outdir(filepath):
    """
    Create a folder to put output files, in the same directory as the input.

    :param filepath: path the the xlsform file being read.
    :return: No return
    """
    outpath = os.path.join(os.path.dirname(filepath), 'out')
    try:
        os.makedirs(outpath)
    except OSError:
        if OSError.errno != errno.EEXIST:
            raise

if __name__ == '__main__':
    # grab the command line arguments
    parser = argparse.ArgumentParser()
    parser.add_argument("-f", "--filepath", help="path to xls form file")
    args = parser.parse_args()
    #create_outdir(args.filepath)
    read_xlsform(args.filepath)
